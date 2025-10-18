import re
from docx import Document
from docx.shared import Inches
from markdown import markdown

def md_to_docx(md_file, docx_file):
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Find the image link
    img_match = re.search(r'!\[(.*?)\]\((.*?)\)', md_content)
    if img_match:
        alt_text = img_match.group(1)
        img_path = img_match.group(2)
        # Remove the image link from the markdown content
        md_content = md_content.replace(img_match.group(0), '')
    else:
        img_path = None

    html_content = markdown(md_content)
    
    document = Document()

    # Add the image if it exists
    if img_path:
        document.add_picture(img_path, width=Inches(6.0))

    # This is a very basic conversion. It will not preserve all formatting.
    # A more robust solution would involve parsing the HTML and converting each element.
    # For this task, we will add the HTML content as a single block of text.
    # A better approach would be to parse the markdown and add elements one by one.
    
    # A simple way to add the content is to split it by lines and add each line as a paragraph.
    # This will lose some formatting but is a step up from adding the whole block.
    lines = md_content.split('\n')
    for line in lines:
        if line.startswith('# '):
            document.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            document.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            document.add_heading(line[4:], level=3)
        else:
            document.add_paragraph(line)

    document.save(docx_file)

if __name__ == '__main__':
    md_to_docx('manuscript.md', 'manuscript.docx')
